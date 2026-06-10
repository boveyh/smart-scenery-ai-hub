"""
景区知识库本地导入脚本
=======================
将 docx / excel 数据集解析 → 文本分块 → 写入本地 JSON 知识库

用途：
  - 后端 ES 入库未就绪时，AI 引擎可直接读取本地 JSON 做开发自测
  - 后续成员B 做完管理后台后，数据通过 ES 接口入库，此脚本可退役

用法：
  cd ai-engine-python
  pip install python-docx openpyxl
  python scripts/import_knowledge.py --tenant west_lake \
      --docx 景区介绍.docx \
      --docx 导览词.docx \
      --excel POI信息.xlsx

输出：
  ai-engine-python/data/knowledge/west_lake.json   （多租户知识库 JSON）
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

# ─── 添加项目根目录到 sys.path ────────────────────────────
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


def parse_docx(file_path: str) -> str:
    """解析 Word 文档，提取全部文本"""
    try:
        from docx import Document
    except ImportError:
        print("❌ 请先安装 python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n".join(paragraphs)


def parse_excel(file_path: str) -> str:
    """解析 Excel 文件，将每行转为一个结构化文本片段"""
    try:
        import openpyxl
    except ImportError:
        print("❌ 请先安装 openpyxl: pip install openpyxl")
        sys.exit(1)

    wb = openpyxl.load_workbook(file_path, data_only=True)
    chunks = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [str(h).strip() if h else f"列{i}" for i, h in enumerate(rows[0])]

        for row in rows[1:]:
            if all(v is None for v in row):
                continue
            parts = []
            for header, value in zip(headers, row):
                if value is not None and str(value).strip():
                    parts.append(f"{header}：{str(value).strip()}")
            if parts:
                chunks.append("，".join(parts))

    return "\n".join(chunks)


def chunk_text(text: str, max_chars: int = 300, overlap: int = 50) -> list[str]:
    """
    文本分块 (Chunking)
    
    策略：
      1. 优先按自然段落分块
      2. 过长段落按句号再次切分
      3. 块间保留 overlap 重叠，避免检索断裂
    """
    if not text:
        return []

    # ─── Step 1: 按自然段落拆分 ──────────────────────────
    raw_paragraphs = re.split(r"\n\s*\n", text)

    # ─── Step 2: 过长段落二次切分 ────────────────────────
    chunks = []
    for para in raw_paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(para) <= max_chars:
            chunks.append(para)
        else:
            # 按句号、问号、感叹号切分
            sentences = re.split(r"(?<=[。！？!?])", para)
            current = ""
            for sent in sentences:
                if len(current) + len(sent) <= max_chars:
                    current += sent
                else:
                    if current.strip():
                        chunks.append(current.strip())
                    # 带重叠的滑动窗口
                    if len(current) > overlap:
                        current = current[-overlap:] + sent
                    else:
                        current = sent
            if current.strip():
                chunks.append(current.strip())

    return chunks


def build_knowledge_base(
    tenant_id: str,
    docx_files: list[str],
    excel_files: list[str],
    output_dir: str = "data/knowledge",
) -> dict:
    """
    构建租户知识库 JSON

    Returns:
        {
            "tenant_id": "west_lake",
            "chunks": [
                {"id": "chunk_001", "content": "...", "source": "景区介绍.docx"},
                ...
            ],
            "metadata": {"total_chunks": 42, "sources": [...]}
        }
    """
    all_chunks = []

    # ─── 处理 Word 文件 ──────────────────────────────────
    for docx_path in docx_files:
        if not os.path.exists(docx_path):
            print(f"⚠️  文件不存在，跳过: {docx_path}")
            continue

        print(f"📄 解析 DOCX: {docx_path}")
        raw_text = parse_docx(docx_path)
        chunks = chunk_text(raw_text)
        source_name = os.path.basename(docx_path)
        for chunk in chunks:
            all_chunks.append({
                "content": chunk,
                "source": source_name,
            })
        print(f"   → 提取 {len(chunks)} 个文本块")

    # ─── 处理 Excel 文件 ────────────────────────────────
    for xlsx_path in excel_files:
        if not os.path.exists(xlsx_path):
            print(f"⚠️  文件不存在，跳过: {xlsx_path}")
            continue

        print(f"📊 解析 EXCEL: {xlsx_path}")
        raw_text = parse_excel(xlsx_path)
        chunks = chunk_text(raw_text)
        source_name = os.path.basename(xlsx_path)
        for chunk in chunks:
            all_chunks.append({
                "content": chunk,
                "source": source_name,
            })
        print(f"   → 提取 {len(chunks)} 个文本块")

    # ─── 写入 JSON ──────────────────────────────────────
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    json_file = output_path / f"{tenant_id}.json"

    knowledge = {
        "tenant_id": tenant_id,
        "chunks": [
            {"id": f"{tenant_id}_chunk_{i:04d}", **chunk}
            for i, chunk in enumerate(all_chunks)
        ],
        "metadata": {
            "total_chunks": len(all_chunks),
            "sources": [os.path.basename(f) for f in docx_files + excel_files],
            "created_at": __import__("datetime").datetime.now().isoformat(),
        },
    }

    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(knowledge, f, ensure_ascii=False, indent=2)

    print(f"\n✅ 知识库已生成: {json_file}")
    print(f"   租户: {tenant_id}")
    print(f"   总文本块: {len(all_chunks)}")
    print(f"   文件大小: {json_file.stat().st_size / 1024:.1f} KB")

    return knowledge


# ─── CLI 入口 ─────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="景区知识库本地导入工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python scripts/import_knowledge.py \\
      --tenant west_lake \\
      --docx "data/西湖景区介绍.docx" \\
      --docx "data/导览词.docx" \\
      --excel "data/POI信息.xlsx"
        """,
    )
    parser.add_argument("--tenant", required=True, help="租户ID，如 west_lake")
    parser.add_argument("--docx", action="append", default=[], help="Word 文件路径（可多次指定）")
    parser.add_argument("--excel", action="append", default=[], help="Excel 文件路径（可多次指定）")
    parser.add_argument("--output", default="data/knowledge", help="输出目录（默认 data/knowledge）")
    parser.add_argument("--chunk-size", type=int, default=300, help="分块最大字符数（默认300）")

    args = parser.parse_args()

    if not args.docx and not args.excel:
        print("❌ 请至少指定一个 --docx 或 --excel 文件")
        parser.print_help()
        sys.exit(1)

    build_knowledge_base(
        tenant_id=args.tenant,
        docx_files=args.docx,
        excel_files=args.excel,
        output_dir=args.output,
    )


if __name__ == "__main__":
    main()